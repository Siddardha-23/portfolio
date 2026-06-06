"""
Resume document renderer — PDF (fpdf2) and DOCX (python-docx) generation.

Template-based rendering with zero AI involvement.
Format: A4, 0.4in margins, Times font, tight spacing.
Section order: Summary → Experience → Projects → Technical Skills → Education.
"""
import io
import re
from datetime import datetime, timezone
from typing import Any, Dict


class ResumeRenderer:
    """Generates PDF and DOCX documents from a structured (tailored) resume dict."""

    # Bump this whenever the rendered OUTPUT changes (margins, fonts, name
    # casing, section layout, bullet rules, etc.). The download endpoint caches
    # rendered files in S3 keyed on the resume *content* hash — but format
    # changes live here in the renderer, not the content, so without a version
    # tag a cache "hit" would keep serving a PDF rendered by the OLD code.
    # Including this in the cache key forces a re-render after any format fix.
    RENDER_VERSION = "2026-06-06.2"

    @staticmethod
    def _contact_display(key: str, value: str) -> tuple:
        """Return (display_text, href_url_or_None) for a contact field."""
        if not value:
            return ("", None)
        if key == "linkedin":
            display = re.sub(r'^https?://', '', value)  # strip protocol
            href = value if value.startswith("http") else f"https://{value}"
            return (display, href)
        if key == "github":
            display = re.sub(r'^https?://', '', value)
            href = value if value.startswith("http") else f"https://{value}"
            return (display, href)
        if key == "portfolio":
            # Strip protocol for display, link to full URL
            display = re.sub(r'^https?://', '', value).rstrip('/')
            href = value if value.startswith("http") else f"https://{value}"
            return (display, href)
        if key == "email":
            return (value, f"mailto:{value}")
        if key == "location":
            return (value, None)  # plain text, no link
        return (value, None)  # phone — no href

    @staticmethod
    def _title_case_name(name: str) -> str:
        """Normalize candidate name to Title Case.

        If a token is entirely upper- or lower-case, capitalize only the first
        letter ("HARSHITH" -> "Harshith"). If it already has mixed case
        (e.g. "McDonald", "deShawn"), leave it untouched so we don't damage
        legitimately styled surnames.
        """
        if not name:
            return name
        out = []
        for tok in name.split():
            if tok.isupper() or tok.islower():
                out.append(tok[:1].upper() + tok[1:].lower())
            else:
                out.append(tok)
        return " ".join(out)

    @staticmethod
    def _sanitize_text(text) -> str:
        """Replace Unicode chars unsupported by built-in PDF fonts with latin-1 equivalents."""
        if not text:
            return ""
        text = str(text)
        replacements = {
            "\u2022": "\xb7",  # bullet • → middle dot · (latin-1 compatible)
            "\u2013": " - ",   # en-dash – → spaced hyphen
            "\u2014": " - ",   # em-dash — → spaced hyphen
            "\u2015": " - ",   # horizontal bar ―
            "\u2212": "-",     # minus sign − (Latin-1 safe single hyphen)
            "\u2018": "'",     # left single quote
            "\u2019": "'",     # right single quote / apostrophe
            "\u201c": '"',     # left double quote
            "\u201d": '"',     # right double quote
            "\u2026": "...",   # ellipsis …
            "\u00a0": " ",     # non-breaking space
            "\u200b": "",      # zero-width space
        }
        for char, repl in replacements.items():
            text = text.replace(char, repl)
        # Collapse duplicate spaces that the spaced-hyphen replacement can
        # introduce when the original already had a space (e.g. "x — y" →
        # "x  -  y" → "x - y").
        text = re.sub(r"  +", " ", text)
        # Fallback: strip any remaining non-latin-1 characters
        return text.encode("latin-1", errors="replace").decode("latin-1")

    # ------------------------------------------------------------------
    # PDF generation (fpdf2) — matches LaTeX resume.cls template
    # ------------------------------------------------------------------

    # Layout constants — loosened from a previous aggressive-density pass.
    # The earlier ultra-tight values produced a "packed" look where bullets
    # crashed into role titles and section headers crowded their content.
    # These values give a clean professional appearance while staying within
    # one page for typical 2-experience + up-to-3-project resumes. The Pass 2
    # distribution + tier fallback (10pt-tight → 9.5pt → 9pt) still handle
    # the densest layouts automatically.
    _MARGIN_LR = 11.0          # ~0.43 inch — wider text column, fewer wraps
    _MARGIN_TOP = 11.5         # slightly under 12mm to add usable height
    _MARGIN_BOTTOM = 11.5      # symmetric with top
    _PAGE_H = 297.0            # A4 height mm
    _AVAIL_H = _PAGE_H - _MARGIN_TOP - _MARGIN_BOTTOM  # ~276.8mm usable

    # Font sizes
    _NAME_SIZE = 19.0          # large centered name
    _CONTACT_SIZE = 10.0       # contact line
    _BODY_SIZE = 11.0          # body base (overridden by adaptive tier)
    _SECTION_TITLE_SIZE = 11.0 # bold UPPERCASE section header
    _COMPANY_SIZE = 11.0       # bold company name
    _JOB_TITLE_SIZE = 11.0     # italic role title — matches company size for a
                               # clean hierarchy (was 10.5pt, looked subordinate)
    _BULLET_INDENT = 5.05      # ~0.199in

    # Line heights
    _LH_BODY = 4.2             # body line height
    _LH_BULLET = 4.0           # bullet line height
    _LH_SKILL = 4.0            # skill row line height
    _LH_CONTACT = 4.5          # contact line height

    _SECTION_HR_WIDTH = 0.16   # thin rule under section header

    # Spacing floors — minimum visual gaps. Each constant fires multiple times
    # on a project-heavy resume, so values are kept conservative to fit one page:
    #   section_gap fires 5× (Summary, Experience, Projects, Skills, Education)
    #   entry_gap fires up to 4× (between experience entries + between projects)
    #   post_header fires 5× (after each section's HR rule)
    _MIN_SECTION_GAP = 1.5     # before section header — clear visual break
    _MIN_ENTRY_GAP = 1.5       # between experiences / between projects
    _MIN_POST_HEADER = 1.0     # after section HR, before first content
    _MIN_HEADER_GAP = 0.5      # after name+contact block
    _MIN_SKILL_GAP = 0.4       # compact separation between skill categories.
                               # Reference LaTeX uses \vspace{-2pt} (tighter
                               # than touching); we keep ours visible but
                               # snug — the previous 0.8mm made skills look
                               # airier than the rest of the resume.
    _BULLET_GAP = 0.6          # gap between bullets within a list — breathing
                               # room so consecutive bullets don't visually merge
    _PRE_BULLET_GAP = 1.2      # gap below italic role title before first bullet
    _EDU_ENTRY_GAP = 1.2       # gap between two education entries
    _JOB_TITLE_CELL_H = 4.0    # italic role-title line height (was 4.5 — the
                               # extra 0.5mm read as dead space under the title)

    # Spacing ceilings — how much Pass 2 may inflate on sparse pages.
    # Sparse pages (1 exp + 1 proj) get a relaxed look; dense pages stay
    # near the floor because Pass 2 distributes any slack proportionally.
    _MAX_SECTION_GAP = 2.8
    _MAX_ENTRY_GAP = 2.4
    _MAX_POST_HEADER = 1.8
    _MAX_HEADER_GAP = 0.8
    _MAX_SKILL_GAP = 0.9

    def _count_spacing_slots(self, tailored: Dict[str, Any]) -> Dict[str, int]:
        """Count the number of each spacing slot type in the resume."""
        section_count = 0
        entry_gaps = 0
        skill_rows = 0

        if tailored.get("summary"):
            section_count += 1
        experience = tailored.get("experience", [])
        if experience:
            section_count += 1
            entry_gaps += max(0, len(experience) - 1)
        projects = tailored.get("projects", [])
        if projects:
            section_count += 1
            entry_gaps += max(0, len(projects) - 1)
        skills = tailored.get("skills", {})
        if skills:
            section_count += 1
            skill_rows = sum(1 for v in skills.values() if v)
        if tailored.get("education"):
            section_count += 1

        return {
            "sections": section_count,
            "entry_gaps": entry_gaps,
            "skill_rows": skill_rows,
        }

    def _render_pdf(
        self,
        tailored: Dict[str, Any],
        *,
        section_gap: float,
        entry_gap: float,
        post_header: float,
        header_gap: float,
        skill_gap: float,
        body_size: float = 11.0,
        lh: float = 4.2,
        lh_s: float = 4.0,
        measure_only: bool = False,
    ) -> "tuple[bytes | None, float]":
        """Core rendering logic matching the LaTeX resume.cls template.

        Returns (pdf_bytes_or_None, content_height_mm).
        If measure_only is True, pdf_bytes is None (saves memory).
        """
        from fpdf import FPDF
        sanitize = self._sanitize_text

        contact = tailored.get("contact", {})
        name = sanitize(contact.get("name", "")).strip() or "Resume"
        name = self._title_case_name(name)
        ml = self._MARGIN_LR
        mt = self._MARGIN_TOP

        pdf = FPDF(format="A4")
        pdf.set_auto_page_break(auto=False)
        # ATS hint: enterprise screeners (Workday, iCIMS) sometimes flag
        # PDFs with empty Document Information as low-confidence — populating
        # Title/Author/Subject also makes the file readable when recruiters
        # save it locally and sort by metadata.
        target_role = sanitize(str(tailored.get("target_role") or "")).strip()
        pdf.set_title(f"Resume - {name}")
        pdf.set_author(name)
        pdf.set_subject(target_role or "Resume")
        pdf.set_creator("Portfolio Resume Builder")
        pdf.add_page()
        pdf.set_margins(ml, mt, ml)
        pdf.set_y(mt)
        w = pdf.w - 2 * ml  # usable width ≈ 184.6mm

        # Name in Title Case (centered, bold). Section headers below still
        # use uppercase, but the candidate's name itself reads more naturally
        # as "Harshith Siddardha Manne" than as ALL CAPS.
        pdf.set_font("Times", "B", self._NAME_SIZE)
        pdf.cell(w, 6.5, name, align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(0.5)  # \nameskip = \smallskip

        # ── Contact line (centered, pipe-separated, black text) ──
        # LaTeX: \fontsize{11}{13}\selectfont, pipe separator
        # Order matches the LaTeX template's reference resume:
        #   location | phone | email | portfolio | linkedin | github
        contact_fields = []
        for key in ("location", "phone", "email", "portfolio", "linkedin", "github"):
            val = contact.get(key, "")
            if val:
                display, href = self._contact_display(key, val)
                contact_fields.append((sanitize(display), href))

        if contact_fields:
            pdf.set_font("Times", "", self._CONTACT_SIZE)
            pdf.set_text_color(0, 0, 0)
            sep = " | "
            sep_w = pdf.get_string_width(sep)

            # Wrap-aware layout: distribute contact fields across multiple
            # lines so the line never exceeds page width. Previously a single
            # too-wide line produced a negative x_start, pushing most of the
            # text off the page (only the right-most fragment showed up).
            # Matches the reference template which wraps when total width
            # exceeds the page (Tempe, AZ | phone | email | portfolio |
            # linkedin |\n github).
            lines: List[List[Tuple[str, Optional[str]]]] = []
            current_line: List[Tuple[str, Optional[str]]] = []
            current_width = 0.0
            for display, href in contact_fields:
                field_w = pdf.get_string_width(display)
                added_w = field_w if not current_line else (sep_w + field_w)
                if current_line and (current_width + added_w) > w:
                    lines.append(current_line)
                    current_line = [(display, href)]
                    current_width = field_w
                else:
                    current_line.append((display, href))
                    current_width += added_w
            if current_line:
                lines.append(current_line)

            for line in lines:
                line_w = sum(pdf.get_string_width(d) for d, _ in line)
                line_w += sep_w * max(0, len(line) - 1)
                x_start = ml + max(0.0, (w - line_w) / 2)
                pdf.set_x(x_start)
                for i, (display, href) in enumerate(line):
                    cell_w = pdf.get_string_width(display)
                    if href:
                        pdf.cell(cell_w, self._LH_CONTACT, display, link=href)
                    else:
                        pdf.cell(cell_w, self._LH_CONTACT, display)
                    if i < len(line) - 1:
                        pdf.cell(sep_w, self._LH_CONTACT, sep)
                pdf.ln(self._LH_CONTACT)

        pdf.ln(header_gap)

        # ── Helper: section header — LaTeX: \sectionskip, UPPERCASE title, \hrule 0.5pt ──
        def section_header(title: str):
            pdf.ln(section_gap)
            pdf.set_font("Times", "B", self._SECTION_TITLE_SIZE)
            pdf.cell(w, 4.5, title.upper(), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(0.5)  # \sectionlineskip = \medskip (small gap before rule)
            ry = pdf.get_y()
            pdf.set_draw_color(0, 0, 0)
            pdf.set_line_width(self._SECTION_HR_WIDTH)
            pdf.line(ml, ry, ml + w, ry)
            pdf.ln(post_header)

        # ── Helper: bullet point — constant spacing between bullets ──
        def bullet(text: str, is_first: bool = False):
            if is_first:
                pdf.ln(self._PRE_BULLET_GAP)  # gap before first bullet
            else:
                pdf.ln(self._BULLET_GAP)      # constant gap between bullets
            indent = self._BULLET_INDENT
            pdf.set_font("Times", "", body_size)
            x_start = pdf.l_margin + indent
            y_center = pdf.get_y() + lh / 2

            # Small filled circle bullet (Latin-1 safe)
            pdf.set_fill_color(0, 0, 0)
            pdf.circle(x_start + 0.8, y_center, 0.6, style="F")

            pdf.set_x(x_start + 2.5)
            pdf.multi_cell(w - indent - 2.5, lh, sanitize(text),
                           new_x="LMARGIN", new_y="NEXT")

        # ── SUMMARY ── (LaTeX: rSection{SUMMARY})
        # Reference template uses the short "SUMMARY" header. First
        # descriptor phrase rendered bold via multi_cell markdown — safer
        # than pdf.write() which doesn't reset X reliably between
        # font-weight switches.
        summary = tailored.get("summary", "")
        if summary:
            section_header("Summary")
            pdf.set_font("Times", "", body_size)
            safe_summary = sanitize(summary)
            # Detect the lead descriptor phrase: everything up to the first
            # connective ("with", "who", "experienced", "skilled",
            # "specializing", "focused", "that", "building", "designing",
            # comma, en/em-dash). Bold only if ≤ 60 chars and ≥ 4 chars.
            lead_match = re.match(
                r'^([^,—\-]+?)(?=\s+(?:with|who|experienced|skilled|specializing|focused|that|building|designing)\b|\s*[—,]\s*|\s+-\s+)',
                safe_summary,
                flags=re.IGNORECASE,
            )
            lead = lead_match.group(1).strip() if lead_match else ""
            if lead and 3 < len(lead) <= 60 and "**" not in safe_summary:
                # Wrap lead in markdown bold; multi_cell handles wrapping
                # within `w` correctly, unlike pdf.write().
                rest = safe_summary[len(lead):]
                rendered = f"**{lead}**{rest}"
                pdf.multi_cell(w, lh, rendered, markdown=True,
                               new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.multi_cell(w, lh, safe_summary,
                               new_x="LMARGIN", new_y="NEXT")

        # ── EXPERIENCE ── (LaTeX: rSection{EXPERIENCE})
        experience = tailored.get("experience", [])
        if experience:
            section_header("Experience")
            for i, exp in enumerate(experience):
                company = sanitize(exp.get("company", ""))
                location = sanitize(exp.get("location", ""))
                # Dates are already normalized by services.date_normalizer
                # before reaching the renderer (e.g. "August 2025 - Present"
                # with ASCII hyphen). Pass through unchanged so the rendered
                # PDF text matches the JSON exactly — the ATS parseability
                # check tokenizes the JSON and requires those tokens to
                # appear verbatim in the extracted PDF text.
                dates = sanitize(exp.get("dates", ""))
                company_line = f"{company}, {location}" if location else company

                # Company bold + dates right — LaTeX: \textbf{COMPANY} \hfill {dates}
                pdf.set_font("Times", "B", self._COMPANY_SIZE)
                if dates:
                    pdf.set_font("Times", "", body_size)
                    date_w = pdf.get_string_width(dates) + 2
                    pdf.set_font("Times", "B", self._COMPANY_SIZE)
                    pdf.cell(w - date_w, 4.5, company_line)
                    pdf.set_font("Times", "", body_size)
                    pdf.cell(date_w, 4.5, dates, align="R",
                             new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.cell(w, 4.5, company_line,
                             new_x="LMARGIN", new_y="NEXT")

                # Job title italic — LaTeX: \fontsize{12}{14}\selectfont\textit{title}
                title = sanitize(exp.get("title", ""))
                if title:
                    pdf.set_font("Times", "I", self._JOB_TITLE_SIZE)
                    pdf.cell(w, self._JOB_TITLE_CELL_H, title,
                             new_x="LMARGIN", new_y="NEXT")

                exp_bullets = exp.get("bullets", [])
                for bi, b in enumerate(exp_bullets):
                    bullet(b, is_first=(bi == 0))
                if i < len(experience) - 1:
                    pdf.ln(entry_gap)

        # ── PROJECTS ── (LaTeX: rSection{PROJECTS})
        projects = tailored.get("projects", [])
        if projects:
            section_header("Projects")
            for i, proj in enumerate(projects):
                # Project name bold — LaTeX: \textbf{Project Name}
                pdf.set_font("Times", "B", self._COMPANY_SIZE)
                proj_name = sanitize(proj.get("name", ""))
                proj_dates = sanitize(proj.get("dates", ""))
                if proj_dates:
                    pdf.set_font("Times", "", body_size)
                    date_w = pdf.get_string_width(proj_dates) + 2
                    pdf.set_font("Times", "B", self._COMPANY_SIZE)
                    pdf.cell(w - date_w, 4.5, proj_name)
                    pdf.set_font("Times", "", body_size)
                    pdf.cell(date_w, 4.5, proj_dates, align="R",
                             new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.cell(w, 4.5, proj_name,
                             new_x="LMARGIN", new_y="NEXT")

                proj_bullets = proj.get("bullets", [])
                for bi, b in enumerate(proj_bullets):
                    bullet(b, is_first=(bi == 0))
                if i < len(projects) - 1:
                    pdf.ln(entry_gap)

        # ── TECHNICAL SKILLS ── (LaTeX: rSection{TECHNICAL SKILLS})
        skills = tailored.get("skills", {})
        if skills:
            section_header("Technical Skills")
            # Use a slightly looser line-height than the bullet tight value
            # so a category that wraps onto a second line still reads as one
            # cohesive block. Without this, wrapped categories collapse onto
            # the following category and the boundary disappears.
            lh_skill = max(lh_s, 3.45)
            for category, skill_list in skills.items():
                if not skill_list:
                    continue
                # Bold category + regular skills in one multi_cell via markdown.
                # multi_cell wraps all lines to the left margin.
                pdf.set_font("Times", "", body_size)
                sk = ", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list)
                line = f"**{sanitize(category)}:** {sanitize(sk)}"
                pdf.multi_cell(w, lh_skill, line, markdown=True,
                               new_x="LMARGIN", new_y="NEXT")
                pdf.ln(skill_gap)

        # ── EDUCATION ── (LaTeX: rSection{Education})
        education = tailored.get("education", [])
        if education:
            seen_edu = set()
            deduped_education = []
            for edu in education:
                inst = sanitize(edu.get("institution", "")).strip()
                degree = sanitize(edu.get("degree", "")).strip()
                if not inst and not degree:
                    continue
                key = (inst.lower(), degree.lower())
                if key in seen_edu:
                    continue
                seen_edu.add(key)
                deduped_education.append(edu)

            if deduped_education:
                section_header("Education")
                for ei, edu in enumerate(deduped_education):
                    if ei > 0:
                        pdf.ln(self._EDU_ENTRY_GAP)  # gap between education entries
                    inst = sanitize(edu.get("institution", ""))
                    location = sanitize(edu.get("location", ""))
                    degree = sanitize(edu.get("degree", ""))
                    dates = sanitize(edu.get("dates", ""))
                    gpa = sanitize(edu.get("gpa", ""))

                    # Strip date patterns baked into degree by AI
                    if degree and re.search(
                        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}\b',
                        degree, re.IGNORECASE,
                    ):
                        cleaned = re.sub(
                            r'\s*\|?\s*\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*'
                            r'\s+\d{4}\b',
                            '', degree, flags=re.IGNORECASE,
                        ).strip().rstrip('|').strip()
                        if cleaned:
                            degree = cleaned

                    # Reference layout:
                    #   Institution[, Location] | Degree[, GPA: X.XX]   <date right-aligned>
                    # Skip location if it's already substring of institution
                    # (e.g. "JNTU Kakinada" + location "Kakinada" → don't
                    # render "JNTU Kakinada, Kakinada").
                    if inst and location and location.lower() in inst.lower():
                        inst_part = inst
                    elif inst and location:
                        inst_part = f"{inst}, {location}"
                    else:
                        inst_part = inst
                    left_parts = [p for p in [inst_part, degree] if p]
                    if gpa:
                        left_parts.append(f"GPA: {gpa}")
                    left_text = " | ".join(left_parts)

                    pdf.set_font("Times", "", body_size)
                    if dates:
                        date_w = pdf.get_string_width(dates) + 2
                        pdf.cell(w - date_w, lh_s, left_text)
                        pdf.cell(date_w, lh_s, dates, align="R",
                                 new_x="LMARGIN", new_y="NEXT")
                    else:
                        pdf.multi_cell(w, lh_s, left_text,
                                       new_x="LMARGIN", new_y="NEXT")

        content_h = pdf.get_y() - mt
        if measure_only:
            return None, content_h
        return bytes(pdf.output()), content_h

    def generate_pdf(self, tailored: Dict[str, Any]) -> bytes:
        """Generate a single-page PDF that adaptively fills one A4 page.

        Two-pass approach:
          Pass 1: Render with minimum spacing to measure content height.
          Pass 2: Distribute remaining whitespace proportionally, then render final PDF.

        If content overflows even at minimum spacing, reduces font size and line heights.
        """
        avail = self._AVAIL_H
        slots = self._count_spacing_slots(tailored)

        # --- Pass 1: measure with minimum spacing at the base render font size (10pt) ---
        # User requested 10pt body to match reference resume's font size.
        # Combined with tighter margins (0.25"/0.4" top/bottom) and lh=3.4
        # (~96% of nominal 11.4pt-at-10pt-body line spacing — comparable to
        # the reference LaTeX template's 0.92 linespread), the candidate's
        # 5/3/3 experience + 5 project bullets + 5-6 skill categories fits
        # at 10pt with a few mm of leftover. Auto-shrink fallback to 9.5pt
        # → 9pt remains for cases where content grows (extra creative
        # project, longer bullets, etc).
        body_size = 10.0
        lh = 3.4
        lh_s = 3.08

        _, min_height = self._render_pdf(
            tailored,
            section_gap=self._MIN_SECTION_GAP,
            entry_gap=self._MIN_ENTRY_GAP,
            post_header=self._MIN_POST_HEADER,
            header_gap=self._MIN_HEADER_GAP,
            skill_gap=self._MIN_SKILL_GAP,
            body_size=body_size, lh=lh, lh_s=lh_s,
            measure_only=True,
        )

        # --- Overflow protection: tiered fallback (10pt-tight → 10pt-ultra → 9.5pt → 9pt) ---
        # Body font size is what recruiters notice most — keep it at 10pt
        # whenever possible. Two tight-lh tiers stack before any font drop:
        #   Tier 1: lh=3.35 / lh_s=3.08 — keeps per-bullet-line breathing room
        #   Tier 2: lh=3.22 / lh_s=3.08
        #   Tier 3: lh=3.10 / lh_s=3.08 — densest 10pt tier before a font drop
        # lh_s (skill/edu rows) is held at 3.08 across all 10pt tiers so those
        # rows keep a consistent line height regardless of which tier fits.
        # Only past Tier 3 do we drop to 9.5pt.
        if min_height > avail:
            for tier_lh, tier_lhs in ((3.35, 3.08), (3.22, 3.08), (3.10, 3.08)):
                _, retry_height = self._render_pdf(
                    tailored,
                    section_gap=self._MIN_SECTION_GAP,
                    entry_gap=self._MIN_ENTRY_GAP,
                    post_header=self._MIN_POST_HEADER,
                    header_gap=self._MIN_HEADER_GAP,
                    skill_gap=self._MIN_SKILL_GAP,
                    body_size=10.0, lh=tier_lh, lh_s=tier_lhs,
                    measure_only=True,
                )
                if retry_height <= avail:
                    lh = tier_lh
                    lh_s = tier_lhs
                    min_height = retry_height
                    break

        if min_height > avail:
            # Still over budget → drop to 9.5pt
            body_size = 9.5
            lh = 3.5
            lh_s = 3.3
            _, min_height = self._render_pdf(
                tailored,
                section_gap=self._MIN_SECTION_GAP,
                entry_gap=self._MIN_ENTRY_GAP,
                post_header=self._MIN_POST_HEADER,
                header_gap=self._MIN_HEADER_GAP,
                skill_gap=self._MIN_SKILL_GAP,
                body_size=body_size, lh=lh, lh_s=lh_s,
                measure_only=True,
            )

        if min_height > avail:
            # Even smaller
            body_size = 9.0
            lh = 3.2
            lh_s = 3.0
            _, min_height = self._render_pdf(
                tailored,
                section_gap=self._MIN_SECTION_GAP,
                entry_gap=self._MIN_ENTRY_GAP,
                post_header=self._MIN_POST_HEADER,
                header_gap=self._MIN_HEADER_GAP,
                skill_gap=self._MIN_SKILL_GAP,
                body_size=body_size, lh=lh, lh_s=lh_s,
                measure_only=True,
            )

        # --- Pass 2: distribute remaining whitespace ---
        remaining = max(0, avail - min_height)

        # Weight distribution: section gaps get 40%, entry gaps 30%,
        # post-header 15%, header gap 10%, skill gaps 5%
        n_sections = max(slots["sections"], 1)
        n_entry_gaps = max(slots["entry_gaps"], 1)
        n_skill_rows = max(slots["skill_rows"], 1)

        total_slots = (
            n_sections * 4      # section_gap weight
            + n_entry_gaps * 3  # entry_gap weight
            + n_sections * 1.5  # post_header weight
            + 1 * 1             # header_gap weight
            + n_skill_rows * 0.5  # skill_gap weight
        )

        if total_slots > 0 and remaining > 0:
            unit = remaining / total_slots
        else:
            unit = 0

        section_gap = min(self._MIN_SECTION_GAP + unit * 4, self._MAX_SECTION_GAP)
        entry_gap = min(self._MIN_ENTRY_GAP + unit * 3, self._MAX_ENTRY_GAP)
        post_header = min(self._MIN_POST_HEADER + unit * 1.5, self._MAX_POST_HEADER)
        header_gap = min(self._MIN_HEADER_GAP + unit * 1, self._MAX_HEADER_GAP)
        skill_gap = min(self._MIN_SKILL_GAP + unit * 0.5, self._MAX_SKILL_GAP)

        # --- Final render ---
        pdf_bytes, _ = self._render_pdf(
            tailored,
            section_gap=section_gap,
            entry_gap=entry_gap,
            post_header=post_header,
            header_gap=header_gap,
            skill_gap=skill_gap,
            body_size=body_size, lh=lh, lh_s=lh_s,
        )
        return pdf_bytes

    # ------------------------------------------------------------------
    # DOCX generation (python-docx) — mirrors the PDF renderer's layout
    # ------------------------------------------------------------------

    # Usable page height in pt (A4 297mm - 0.4in*2 margins ≈ 277mm ≈ 785pt).
    _DOCX_AVAIL_PT = 785.0

    def _estimate_docx_height_pt(
        self,
        tailored: Dict[str, Any],
        *,
        body_size: float,
        line_spacing: float,
        section_gap: float,
        entry_gap: float,
        skill_gap: float,
        header_gap: float,
    ) -> float:
        """Heuristic height estimator in points.

        python-docx can't render, so we estimate: each paragraph contributes
        (line_count * line_spacing) + paragraph_before/after spacing. Line count
        is derived by dividing text length by a characters-per-line constant
        based on body_size (smaller font ⇒ more chars/line).
        """
        # Usable width in mm * (chars per mm at given pt). Times at 10pt
        # averages ~2.1 chars/mm in our 197mm-wide text box → ~414 chars/line.
        # Scale inversely with font size.
        scale = 10.0 / max(body_size, 6.0)
        chars_per_line = 95 * scale  # rough but stable
        h = 0.0

        def para(text: str, size: float = None, bullet: bool = False) -> float:
            s = size or body_size
            line_h = s * 1.15  # ~1.15 line spacing equivalent
            if not text:
                return line_h + 2
            length = len(text) + (3 if bullet else 0)
            lines = max(1, int(length / chars_per_line) + (1 if length % chars_per_line else 0))
            return lines * line_h + 1.5

        # Name + contact + HR
        h += 17 * 1.15 + 2        # name
        contact = tailored.get("contact", {})
        if any(contact.get(k) for k in ("location", "phone", "email", "portfolio", "linkedin", "github")):
            h += 10.5 * 1.15 + 1  # contact line
        h += 4                     # HR line
        h += header_gap

        def section(title: str) -> float:
            return 11 * 1.15 + 2 + section_gap

        def header_block(left: str, dates: str) -> float:
            return max(11 * 1.15, body_size * 1.15) + 2

        if tailored.get("summary"):
            h += section("SUMMARY")
            h += para(tailored["summary"])

        experience = tailored.get("experience", []) or []
        if experience:
            h += section("EXPERIENCE")
            for i, exp in enumerate(experience):
                h += header_block(exp.get("company", ""), exp.get("dates", ""))
                if exp.get("title"):
                    h += 10.5 * 1.15 + 1
                for b in exp.get("bullets", []) or []:
                    h += para(b, bullet=True)
                if i < len(experience) - 1:
                    h += entry_gap

        projects = tailored.get("projects", []) or []
        if projects:
            h += section("PROJECTS")
            for i, proj in enumerate(projects):
                h += header_block(proj.get("name", ""), proj.get("dates", ""))
                for b in proj.get("bullets", []) or []:
                    h += para(b, bullet=True)
                if i < len(projects) - 1:
                    h += entry_gap

        skills = tailored.get("skills", {}) or {}
        skill_rows = [v for v in skills.values() if v]
        if skill_rows:
            h += section("TECHNICAL SKILLS")
            for sk in skill_rows:
                text = ", ".join(sk) if isinstance(sk, list) else str(sk)
                h += para(text) + skill_gap

        education = tailored.get("education", []) or []
        if education:
            h += section("EDUCATION")
            for edu in education:
                inst = edu.get("institution", "")
                degree = edu.get("degree", "")
                if inst or degree:
                    h += para(f"{inst} | {degree}")

        return h

    def _render_docx(
        self,
        tailored: Dict[str, Any],
        *,
        body_size: float,
        title_size: float,
        header_size: float,
        section_gap: float,
        entry_gap: float,
        skill_gap: float,
        header_gap: float,
    ) -> bytes:
        """Render a DOCX mirroring the PDF layout with the given spacing knobs."""
        from docx import Document
        from docx.shared import Pt, Inches, Mm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        FONT = "Times New Roman"
        doc = Document()

        # A4 page, 0.5in L/R margins (match PDF _MARGIN_LR = 12.7mm), 0.3in top.
        for section in doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Normal style
        style = doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(body_size)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.1

        # Usable width in EMU for right-aligned tab stops (page_width - L - R)
        usable_w_emu = section.page_width - section.left_margin - section.right_margin

        # ------------------- helpers -------------------
        def set_run_font(run, size_pt: float, bold=False, italic=False, color="000000"):
            run.font.name = FONT
            run.font.size = Pt(size_pt)
            run.bold = bold
            run.italic = italic
            # Ensure Times font for East Asian fallback (prevents font swap in Word)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rFonts.set(qn(attr), FONT)
            c = OxmlElement("w:color")
            c.set(qn("w:val"), color)
            rPr.append(c)

        def tight(p, before=0.0, after=0.0):
            p.paragraph_format.space_before = Pt(before)
            p.paragraph_format.space_after = Pt(after)
            p.paragraph_format.line_spacing = 1.1

        def add_bottom_border(p, size="6"):
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), size)
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "000000")
            pBdr.append(bottom)
            pPr.append(pBdr)

        def add_section_header(title: str):
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            set_run_font(run, header_size, bold=True)
            tight(p, before=section_gap, after=1.0)
            add_bottom_border(p, size="4")

        def add_bullet(text: str, indent_emu=None):
            p = doc.add_paragraph()
            # Manual bullet: filled circle + space, with hanging indent. Using a
            # raw paragraph (not List Bullet style) avoids Word's auto-numbering
            # XML which can drift across versions.
            r_dot = p.add_run("\u2022  ")
            set_run_font(r_dot, body_size, bold=True)
            r = p.add_run(text)
            set_run_font(r, body_size)
            # Indent bullet glyph + hang text under it.
            p.paragraph_format.left_indent = Inches(0.32)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            tight(p, before=0.5, after=0.3)

        # ------------------- HEADER -------------------
        contact = tailored.get("contact", {})
        name = (contact.get("name", "") or "Resume").strip()
        name = self._title_case_name(name)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        set_run_font(run, title_size, bold=True)
        tight(p, before=0, after=1.0)

        contact_fields = []
        for key in ("location", "phone", "email", "portfolio", "linkedin", "github"):
            val = contact.get(key, "")
            if val:
                display, href = self._contact_display(key, val)
                contact_fields.append((display, href))
        if contact_fields:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, (display, href) in enumerate(contact_fields):
                if href:
                    self._add_docx_hyperlink(p, display, href, FONT, Pt(body_size + 0.5))
                else:
                    run = p.add_run(display)
                    set_run_font(run, body_size + 0.5)
                if i < len(contact_fields) - 1:
                    sep = p.add_run(" | ")
                    set_run_font(sep, body_size + 0.5)
            tight(p, before=0, after=header_gap)

        # ------------------- SUMMARY -------------------
        summary = tailored.get("summary", "")
        if summary:
            add_section_header("Professional Summary")
            p = doc.add_paragraph()
            run = p.add_run(summary)
            set_run_font(run, body_size)
            tight(p, before=1.0, after=0.5)

        # ------------------- helper for a company/dates header line -------------------
        def header_line(left_bold: str, dates: str, bold_size: float):
            p = doc.add_paragraph()
            if dates:
                p.paragraph_format.tab_stops.add_tab_stop(
                    usable_w_emu, WD_TAB_ALIGNMENT.RIGHT)
            r1 = p.add_run(left_bold)
            set_run_font(r1, bold_size, bold=True)
            if dates:
                r2 = p.add_run(f"\t{dates}")
                set_run_font(r2, body_size)
            tight(p, before=1.5, after=0.0)
            return p

        # ------------------- EXPERIENCE -------------------
        experience = tailored.get("experience", []) or []
        if experience:
            add_section_header("Experience")
            for i, exp in enumerate(experience):
                company = exp.get("company", "")
                location = exp.get("location", "")
                dates = exp.get("dates", "")
                company_line = f"{company}, {location}" if location else company
                header_line(company_line, dates, bold_size=body_size + 1)

                title = exp.get("title", "")
                if title:
                    p2 = doc.add_paragraph()
                    r = p2.add_run(title)
                    set_run_font(r, body_size + 0.5, italic=True)
                    tight(p2, before=0, after=0.5)

                for b in exp.get("bullets", []) or []:
                    add_bullet(b)
                if i < len(experience) - 1:
                    # Gap between experience entries via an empty spacer para
                    spacer = doc.add_paragraph()
                    tight(spacer, before=entry_gap, after=0)

        # ------------------- PROJECTS -------------------
        projects = tailored.get("projects", []) or []
        if projects:
            add_section_header("Projects")
            for i, proj in enumerate(projects):
                name = proj.get("name", "")
                dates = proj.get("dates", "")
                header_line(name, dates, bold_size=body_size + 1)

                tech = proj.get("tech", "")
                if tech:
                    p2 = doc.add_paragraph()
                    r = p2.add_run(tech)
                    set_run_font(r, body_size, italic=True, color="555555")
                    tight(p2, before=0, after=0.5)

                for b in proj.get("bullets", []) or []:
                    add_bullet(b)
                if i < len(projects) - 1:
                    spacer = doc.add_paragraph()
                    tight(spacer, before=entry_gap, after=0)

        # ------------------- TECHNICAL SKILLS -------------------
        skills = tailored.get("skills", {}) or {}
        if any(v for v in skills.values()):
            add_section_header("Technical Skills")
            for category, skill_list in skills.items():
                if not skill_list:
                    continue
                p = doc.add_paragraph()
                r_cat = p.add_run(f"{category}: ")
                set_run_font(r_cat, body_size, bold=True)
                sk = ", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list)
                r_sk = p.add_run(sk)
                set_run_font(r_sk, body_size)
                tight(p, before=skill_gap, after=0)

        # ------------------- EDUCATION -------------------
        education = tailored.get("education", []) or []
        if education:
            seen_edu = set()
            deduped_education = []
            for edu in education:
                inst = (edu.get("institution", "") or "").strip()
                degree = (edu.get("degree", "") or "").strip()
                if not inst and not degree:
                    continue
                key = (inst.lower(), degree.lower())
                if key in seen_edu:
                    continue
                seen_edu.add(key)
                deduped_education.append(edu)

            if deduped_education:
                add_section_header("Education")
                for edu in deduped_education:
                    inst = edu.get("institution", "") or ""
                    location = edu.get("location", "") or ""
                    degree = edu.get("degree", "") or ""
                    dates = edu.get("dates", "") or ""
                    gpa = edu.get("gpa", "") or ""

                    if degree and re.search(
                        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}\b',
                        degree, re.IGNORECASE,
                    ):
                        cleaned = re.sub(
                            r'\s*\|?\s*\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*'
                            r'\s+\d{4}\b',
                            '', degree, flags=re.IGNORECASE,
                        ).strip().rstrip('|').strip()
                        if cleaned:
                            degree = cleaned

                    inst_part = f"{inst}, {location}" if inst and location else inst
                    parts = [x for x in [inst_part, degree] if x]
                    if gpa:
                        parts.append(f"GPA: {gpa}")
                    left_text = " | ".join(parts)

                    p = doc.add_paragraph()
                    if dates:
                        p.paragraph_format.tab_stops.add_tab_stop(
                            usable_w_emu, WD_TAB_ALIGNMENT.RIGHT)
                    r = p.add_run(left_text)
                    set_run_font(r, body_size)
                    if dates:
                        r2 = p.add_run(f"\t{dates}")
                        set_run_font(r2, body_size)
                    tight(p, before=1.0, after=0)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def generate_docx(self, tailored: Dict[str, Any]) -> bytes:
        """Generate a single-page DOCX that adaptively fills an A4 page.

        Mirrors the PDF renderer's two-pass strategy:
          Pass 1: estimate content height at body=10pt with minimum gaps.
          Pass 2: shrink font to 9.5/9pt if estimated overflow; then distribute
                  remaining whitespace across section/entry/skill gaps so the
                  page fills naturally (same visual density as the PDF).
        """
        avail = self._DOCX_AVAIL_PT

        # Minimum spacing knobs (pt)
        min_section_gap = 2.0
        min_entry_gap = 2.0
        min_skill_gap = 0.0
        min_header_gap = 2.0

        # Max ceilings so fills don't explode on short resumes
        max_section_gap = 12.0
        max_entry_gap = 10.0
        max_skill_gap = 4.0
        max_header_gap = 10.0

        # Try decreasing font sizes until content fits
        for body_size, title_size, header_size in (
            (10.0, 17.0, 11.0),
            (9.5, 16.5, 10.5),
            (9.0, 16.0, 10.0),
        ):
            min_h = self._estimate_docx_height_pt(
                tailored,
                body_size=body_size,
                line_spacing=1.1,
                section_gap=min_section_gap,
                entry_gap=min_entry_gap,
                skill_gap=min_skill_gap,
                header_gap=min_header_gap,
            )
            if min_h <= avail:
                break

        # Count slots for fill distribution
        n_sections = 0
        n_entry_gaps = 0
        n_skill_rows = 0
        if tailored.get("summary"):
            n_sections += 1
        exp = tailored.get("experience", []) or []
        if exp:
            n_sections += 1
            n_entry_gaps += max(0, len(exp) - 1)
        proj = tailored.get("projects", []) or []
        if proj:
            n_sections += 1
            n_entry_gaps += max(0, len(proj) - 1)
        skills = tailored.get("skills", {}) or {}
        skill_rows = [v for v in skills.values() if v]
        if skill_rows:
            n_sections += 1
            n_skill_rows = len(skill_rows)
        if tailored.get("education"):
            n_sections += 1

        # Distribute remaining whitespace (weights mirror the PDF renderer's)
        remaining = max(0, avail - min_h)
        total_weight = (
            max(n_sections, 1) * 4
            + max(n_entry_gaps, 1) * 3
            + 1 * 1            # header gap
            + max(n_skill_rows, 1) * 0.5
        )
        unit = (remaining / total_weight) if total_weight > 0 and remaining > 0 else 0

        section_gap = min(min_section_gap + unit * 4, max_section_gap)
        entry_gap = min(min_entry_gap + unit * 3, max_entry_gap)
        skill_gap = min(min_skill_gap + unit * 0.5, max_skill_gap)
        header_gap = min(min_header_gap + unit * 1, max_header_gap)

        return self._render_docx(
            tailored,
            body_size=body_size,
            title_size=title_size,
            header_size=header_size,
            section_gap=section_gap,
            entry_gap=entry_gap,
            skill_gap=skill_gap,
            header_gap=header_gap,
        )

    # ------------------------------------------------------------------
    # DOCX hyperlink helper
    # ------------------------------------------------------------------

    @staticmethod
    def _add_docx_hyperlink(paragraph, text: str, url: str, font_name: str, font_size):
        """Add a clickable hyperlink run to a DOCX paragraph using OPC XML."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Create the w:hyperlink element
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create the run inside the hyperlink
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Font name
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rPr.append(rFonts)

        # Font size
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(font_size.pt * 2))  # half-points
        rPr.append(sz)

        # Black color (not default blue)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        rPr.append(color)

        new_run.append(rPr)

        # Text node
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    # ------------------------------------------------------------------
    # Filename helper
    # ------------------------------------------------------------------

    @staticmethod
    def build_filename(tailored: Dict[str, Any], jd_analysis: Dict[str, Any], ext: str) -> str:
        """Build ATS-friendly filename: Firstname_Lastname_JobTitle.ext"""
        contact = tailored.get("contact", {})
        name = contact.get("name", "").strip()
        name = ResumeRenderer._title_case_name(name)
        parts = name.split()
        first = parts[0] if parts else "Resume"
        last = parts[-1] if len(parts) > 1 else ""

        job_title = jd_analysis.get("job_title", "")

        def clean(s: str) -> str:
            s = re.sub(r"[^a-zA-Z0-9 ]+", " ", s).strip()
            s = re.sub(r"\s+", "_", s).strip("_")
            return s

        segments = [clean(first), clean(last), clean(job_title)]
        segments = [s for s in segments if s]
        return "_".join(segments) + f".{ext}"

    @staticmethod
    def build_cover_letter_filename(
        candidate_name: str, job_title: str, company: str, ext: str = "pdf",
    ) -> str:
        """Build ATS-friendly cover letter filename:
            Firstname_Lastname_JobTitle_Cover_Letter.ext

        Mirrors the resume filename convention so recruiters see paired
        documents in a single sorted view. Company is intentionally NOT
        included — many ATS portals reject filenames containing punctuation
        from real company names (e.g. "Two Sigma, LLC"), and the job_title
        already disambiguates the role.
        """
        normalized = ResumeRenderer._title_case_name((candidate_name or "").strip())
        parts = normalized.split()
        first = parts[0] if parts else "Applicant"
        last = parts[-1] if len(parts) > 1 else ""

        def clean(s: str) -> str:
            s = re.sub(r"[^a-zA-Z0-9 ]+", " ", s).strip()
            s = re.sub(r"\s+", "_", s).strip("_")
            return s

        segments = [clean(first), clean(last), clean(job_title), "Cover_Letter"]
        segments = [s for s in segments if s]
        return "_".join(segments) + f".{ext}"
