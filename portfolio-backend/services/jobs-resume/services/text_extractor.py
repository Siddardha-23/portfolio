"""
Local text + hyperlink extraction for resumes (PDF and DOCX).

Purpose
-------
The Gemini multi-modal call used to be asked for a verbatim transcription
(`raw_text`) of the document plus a list of every hyperlink. That output was
the dominant cost in the parse pipeline (3-5k tokens of output → 60-80s of
wall time at Flash output speeds).

For the >95% of resumes that are text-based (not scanned images), we can
extract text and links **locally** with byte-perfect accuracy:

  - PDF: pypdf reads the actual text content stream, and pulls hyperlinks
    from /Annots /A /URI dictionaries. No OCR, no inference.
  - DOCX: python-docx reads body paragraphs + tables + headers + footers +
    text boxes, plus walks the rels of the main part AND every header/footer
    part for hyperlinks (rels are scoped per-part in OOXML).

This module exposes a single boundary — `extract_text()` — that callers can
use without caring about file type. It never raises: a failed extraction
returns ("", []) so the caller can fall back to asking Gemini.

Failure budget
--------------
Local extraction is *advisory*. If anything looks off — encrypted PDF,
broken font encoding producing replacement chars, image-only PDF, malformed
DOCX, wildly unusual layout — `is_text_healthy()` returns False and the
caller falls back to the existing FULL-prompt Gemini path. Behavior in that
case is identical to today's pipeline. We never *worsen* the floor.

Lambda safety
-------------
Both libraries (pypdf, python-docx) are pure Python. No system deps.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Iterable, List, Tuple

logger = logging.getLogger(__name__)

# A document is considered to have a usable text layer if extracted text
# clears this many characters AND has a sane fraction of alphanumerics AND
# is not dominated by replacement / non-decodable glyphs.
_MIN_TEXT_CHARS = 50
_MIN_ALNUM_RATIO = 0.30
_MAX_REPLACEMENT_RATIO = 0.05

# Cap pages we process — defends against pathological multi-page PDFs that
# occasionally slip past the upload size cap. A real resume virtually never
# exceeds this.
_MAX_PDF_PAGES = 15

_HYPERLINK_RELTYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

# OOXML namespace map for XPath queries into DOCX bodies.
_W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# URI schemes we accept as hyperlinks. file:// and javascript: are dropped.
_ACCEPTED_URL_SCHEMES = ("http://", "https://", "mailto:", "tel:")


def is_text_healthy(text: str) -> bool:
    """Return True if locally-extracted text is rich enough to skip LLM transcription.

    Guards against:
      - Empty / near-empty extraction (scanned image PDFs)
      - Garbled binary output from broken font encoding
        (low alphanumeric ratio)
      - Replacement-char dominated output (CID-only fonts without ToUnicode)
    """
    if not text or len(text) < _MIN_TEXT_CHARS:
        return False
    n = len(text)
    alnum = sum(1 for c in text if c.isalnum())
    if (alnum / n) < _MIN_ALNUM_RATIO:
        return False
    # U+FFFD = replacement char, emitted when a glyph has no ToUnicode mapping.
    replacement = text.count("�")
    if (replacement / n) > _MAX_REPLACEMENT_RATIO:
        return False
    return True


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _collect_pdf_urls(annots: Iterable, urls: List[str], seen: set) -> None:
    """Walk a page's /Annots list and append unique /Link → /URI targets."""
    for annot_ref in annots or []:
        try:
            annot = annot_ref.get_object() if hasattr(annot_ref, "get_object") else annot_ref
            if not annot:
                continue
            action = annot.get("/A")
            if not action:
                continue
            # /A may be an indirect object reference
            if hasattr(action, "get_object"):
                action = action.get_object()
            uri = action.get("/URI") if action else None
            if not uri:
                continue
            uri_str = str(uri).strip()
            if not uri_str:
                continue
            lower = uri_str.lower()
            # Drop non-resume URI schemes (javascript, file, embedded launches)
            if not any(lower.startswith(s) for s in _ACCEPTED_URL_SCHEMES):
                continue
            if uri_str not in seen:
                seen.add(uri_str)
                urls.append(uri_str)
        except Exception:
            continue


def extract_pdf_text(file_bytes: bytes) -> Tuple[str, List[str]]:
    """Extract text + hyperlink URIs from a PDF.

    Uses pypdf's text-content-stream extraction (no OCR). Pulls hyperlinks
    from each page's /Annots dictionary so we get the actual href targets,
    not display text.

    Robustness:
      - Encrypted PDFs: try empty-password decrypt; if it fails, return
        empty (caller falls back to Gemini OCR).
      - Page cap: process at most _MAX_PDF_PAGES pages.
      - Per-page failures swallowed so one bad page doesn't lose the rest.

    Returns:
        (raw_text, urls). Empty tuple on unrecoverable failure.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf not installed — local PDF extraction unavailable")
        return "", []

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        logger.warning("PDF parse failed in pypdf: %s", e)
        return "", []

    # Encrypted PDFs: try empty password first (common for "owner-locked but
    # readable" exports). If that fails, give up and let the LLM path handle it.
    if getattr(reader, "is_encrypted", False):
        try:
            ok = reader.decrypt("")
            if not ok:
                logger.info("Encrypted PDF (non-empty password) — local extraction skipped")
                return "", []
        except Exception as e:
            logger.info("Encrypted PDF, decrypt failed (%s) — local extraction skipped", e)
            return "", []

    pages_text: List[str] = []
    urls: List[str] = []
    seen_urls: set = set()

    pages = list(reader.pages)[:_MAX_PDF_PAGES]
    for page_num, page in enumerate(pages):
        # ---- text layer ----
        # Default extraction preserves intra-word spacing reliably across the
        # PDFs we tested. Layout-mode is geometrically more accurate for
        # cross-column gaps but, on Word/LaTeX-generated resumes, often loses
        # spaces *inside* words ("HARSHITHSIDDARDHAMANNE") which corrupts the
        # name regex more than cross-column gluing affects anything else.
        try:
            txt = page.extract_text() or ""
        except Exception as e:
            logger.debug("pypdf text extraction failed on page %d: %s", page_num, e)
            txt = ""
        if txt:
            pages_text.append(txt)

        # ---- hyperlink annotations ----
        try:
            annots = page.get("/Annots")
            if annots:
                _collect_pdf_urls(annots, urls, seen_urls)
        except Exception as e:
            logger.debug("Annot iteration failed on page %d: %s", page_num, e)

    return ("\n".join(pages_text), urls)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _walk_paragraphs(container, lines: List[str]) -> None:
    """Append non-empty paragraph text from any container that exposes `.paragraphs`."""
    for para in getattr(container, "paragraphs", []) or []:
        text = (para.text or "").strip()
        if text:
            lines.append(text)


def _walk_tables(tables, lines: List[str]) -> None:
    """Walk every cell of every table, descending into nested tables."""
    for table in tables or []:
        for row in table.rows:
            for cell in row.cells:
                _walk_paragraphs(cell, lines)
                # Some templates nest tables inside cells (calendar grids etc.)
                nested = getattr(cell, "tables", None)
                if nested:
                    _walk_tables(nested, lines)


def _extract_docx_textbox_text(doc) -> List[str]:
    """Pull text out of OOXML text boxes / shapes.

    python-docx doesn't expose drawings/shapes, but every visible run of text
    in OOXML is a `<w:t>` element somewhere in the part's XML. Querying the
    body for w:t nodes that the paragraph walker DIDN'T already reach (i.e.
    those nested in a w:txbxContent) recovers text-box content used by some
    Word templates.
    """
    try:
        from lxml import etree  # noqa: F401  -- only need the etree dep python-docx already pulls in
    except ImportError:
        return []

    out: List[str] = []
    try:
        # Find every textbox content node, then collect descendant w:t runs.
        body = doc.element.body
        for txbx in body.iter("{%s}txbxContent" % _W_NS["w"]):
            for t in txbx.iter("{%s}t" % _W_NS["w"]):
                if t.text:
                    cleaned = t.text.strip()
                    if cleaned:
                        out.append(cleaned)
    except Exception as e:
        logger.debug("DOCX textbox walk failed: %s", e)
    return out


def _walk_section_part(part, lines: List[str], urls: List[str], seen: set) -> None:
    """Walk a header/footer part's paragraphs, tables, and hyperlink rels."""
    if part is None:
        return
    _walk_paragraphs(part, lines)
    _walk_tables(getattr(part, "tables", []) or [], lines)
    try:
        rels = getattr(part, "part", None)
        rels = rels.rels if rels is not None else None
        if rels:
            for rel in rels.values():
                if rel.reltype == _HYPERLINK_RELTYPE:
                    target = (rel.target_ref or "").strip()
                    if target and target not in seen:
                        lower = target.lower()
                        if any(lower.startswith(s) for s in _ACCEPTED_URL_SCHEMES) or "://" not in lower:
                            # Some DOCX exports drop the scheme; keep relative-looking
                            # values too (the consuming regexes can deal with them).
                            seen.add(target)
                            urls.append(target)
    except Exception as e:
        logger.debug("DOCX header/footer rel walk failed: %s", e)


def extract_docx_text(file_bytes: bytes) -> Tuple[str, List[str]]:
    """Extract text + hyperlink URLs from a DOCX file.

    Walks: body paragraphs → body tables (with nested tables) → text boxes
    via OOXML XPath → headers and footers (paragraphs + tables + hyperlink
    rels). Each header/footer part has its own rel collection in OOXML, so
    we walk them separately rather than relying solely on the main part.
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed — local DOCX extraction unavailable")
        return "", []

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.warning("DOCX parse failed in python-docx: %s", e)
        return "", []

    lines: List[str] = []
    urls: List[str] = []
    seen: set = set()

    # ---- body ----
    _walk_paragraphs(doc, lines)
    _walk_tables(doc.tables, lines)
    lines.extend(_extract_docx_textbox_text(doc))

    # ---- headers & footers (per-section) ----
    try:
        for section in doc.sections:
            for hdr_attr in ("header", "first_page_header", "even_page_header"):
                _walk_section_part(getattr(section, hdr_attr, None), lines, urls, seen)
            for ftr_attr in ("footer", "first_page_footer", "even_page_footer"):
                _walk_section_part(getattr(section, ftr_attr, None), lines, urls, seen)
    except Exception as e:
        logger.debug("DOCX section walk failed: %s", e)

    # ---- main-part hyperlinks ----
    try:
        for rel in doc.part.rels.values():
            if rel.reltype == _HYPERLINK_RELTYPE:
                target = (rel.target_ref or "").strip()
                if target and target not in seen:
                    seen.add(target)
                    urls.append(target)
    except Exception as e:
        logger.debug("DOCX main-part rel walk failed: %s", e)

    return ("\n".join(lines), urls)


# ---------------------------------------------------------------------------
# Dispatch + cleaning
# ---------------------------------------------------------------------------


def extract_text(file_bytes: bytes, mime_type: str) -> Tuple[str, List[str]]:
    """Dispatch to the right extractor based on MIME type.

    Returns ("", []) for unsupported types or any failure — callers should
    treat that as "fall back to LLM transcription".
    """
    if not file_bytes:
        return "", []
    mt = (mime_type or "").lower()
    if "pdf" in mt:
        return extract_pdf_text(file_bytes)
    if "word" in mt or "officedocument" in mt or "docx" in mt:
        return extract_docx_text(file_bytes)
    return "", []


# Marker format kept in sync with ResumeParser._backfill_contact, which greps
# for "[Extracted Link: <url>]" tokens to source authoritative hyperlinks.
def build_link_markers(urls: List[str]) -> str:
    if not urls:
        return ""
    return "\n".join(f"[Extracted Link: {u}]" for u in urls if u)


# Strip control characters that pypdf occasionally leaves behind
_CTRL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
# Zero-width / formatting glyphs and BOM. NOTE: U+FFFD is NOT stripped here —
# is_text_healthy treats high U+FFFD density as "fall back to Gemini", but a
# few replacement chars in otherwise-good text (e.g. for one un-mappable
# bullet) shouldn't be silently rewritten away.
_ZW_RE = re.compile(r"[​‌‍﻿]")
# Collapse runs of 3+ blank lines.
_BLANK_LINES_RE = re.compile(r"\n{3,}")


def clean_extracted_text(text: str) -> str:
    if not text:
        return ""
    text = _CTRL_RE.sub("", text)
    text = _ZW_RE.sub("", text)
    text = _BLANK_LINES_RE.sub("\n\n", text)
    return text.strip()
