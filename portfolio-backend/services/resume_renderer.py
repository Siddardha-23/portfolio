"""
Resume document renderer — PDF (fpdf2) and DOCX (python-docx) generation.

Template-based rendering with zero AI involvement.
Format: A4, 0.4in margins, Times font, tight spacing.
Section order: Summary → Experience → Projects → Technical Skills → Education.
"""
import io
import re
from datetime import datetime, timezone
from typing import Any, Dict


class ResumeRenderer:
    """Generates PDF and DOCX documents from a structured (tailored) resume dict."""

    @staticmethod
    def _sanitize_text(text) -> str:
        """Replace Unicode chars unsupported by built-in PDF fonts with latin-1 equivalents."""
        if not text:
            return ""
        text = str(text)
        replacements = {
            "\u2022": "-",   # bullet •
            "\u2013": "-",   # en-dash –
            "\u2014": "--",  # em-dash —
            "\u2018": "'",   # left single quote
            "\u2019": "'",   # right single quote / apostrophe
            "\u201c": '"',   # left double quote
            "\u201d": '"',   # right double quote
            "\u2026": "...", # ellipsis …
            "\u00a0": " ",   # non-breaking space
            "\u200b": "",    # zero-width space
        }
        for char, repl in replacements.items():
            text = text.replace(char, repl)
        # Fallback: strip any remaining non-latin-1 characters
        return text.encode("latin-1", errors="replace").decode("latin-1")

    # ------------------------------------------------------------------
    # PDF generation (fpdf2) — adaptive two-pass one-page renderer
    # ------------------------------------------------------------------

    # Layout constants
    _MARGIN = 10.16            # 0.4 inch in mm
    _PAGE_H = 297.0            # A4 height mm
    _BOTTOM_MARGIN = 10.0      # bottom margin mm
    _AVAIL_H = _PAGE_H - _MARGIN - _BOTTOM_MARGIN  # ~276.8mm usable

    # Spacing defaults (minimum values used during measurement pass)
    _MIN_SECTION_GAP = 1.5     # before section header
    _MIN_ENTRY_GAP = 1.0       # between experience/project entries
    _MIN_POST_HEADER = 1.0     # after section header rule
    _MIN_HEADER_GAP = 1.5      # after name+contact HR
    _MIN_SKILL_GAP = 0.2       # between skill rows

    # Spacing ceilings (prevent over-expansion on very short resumes)
    _MAX_SECTION_GAP = 8.0
    _MAX_ENTRY_GAP = 5.0
    _MAX_POST_HEADER = 4.0
    _MAX_HEADER_GAP = 5.0
    _MAX_SKILL_GAP = 2.0

    def _count_spacing_slots(self, tailored: Dict[str, Any]) -> Dict[str, int]:
        """Count the number of each spacing slot type in the resume."""
        section_count = 0
        entry_gaps = 0
        skill_rows = 0

        if tailored.get("summary"):
            section_count += 1
        experience = tailored.get("experience", [])
        if experience:
            section_count += 1
            entry_gaps += max(0, len(experience) - 1)
        projects = tailored.get("projects", [])
        if projects:
            section_count += 1
            entry_gaps += max(0, len(projects) - 1)
        skills = tailored.get("skills", {})
        if skills:
            section_count += 1
            skill_rows = sum(1 for v in skills.values() if v)
        if tailored.get("education"):
            section_count += 1

        return {
            "sections": section_count,
            "entry_gaps": entry_gaps,
            "skill_rows": skill_rows,
        }

    def _render_pdf(
        self,
        tailored: Dict[str, Any],
        *,
        section_gap: float,
        entry_gap: float,
        post_header: float,
        header_gap: float,
        skill_gap: float,
        body_size: float = 10.0,
        lh: float = 3.8,
        lh_s: float = 3.6,
        measure_only: bool = False,
    ) -> "tuple[bytes | None, float]":
        """Core rendering logic with parameterized spacing.

        Returns (pdf_bytes_or_None, content_height_mm).
        If measure_only is True, pdf_bytes is None (saves memory).
        """
        from fpdf import FPDF
        sanitize = self._sanitize_text

        contact = tailored.get("contact", {})
        name = sanitize(contact.get("name", "Resume"))
        m = self._MARGIN

        pdf = FPDF(format="A4")
        pdf.set_auto_page_break(auto=False)  # We control page overflow
        pdf.add_page()
        pdf.set_margins(m, m, m)
        pdf.set_y(m)
        w = pdf.w - 2 * m  # usable width ~189.7mm

        # ── Name (centered, bold, large) ──
        pdf.set_font("Times", "B", 22)
        pdf.cell(w, 7, name, align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(0.3)

        # ── Contact line ──
        parts = []
        for key in ("phone", "email", "linkedin", "github"):
            val = contact.get(key, "")
            if val:
                parts.append(sanitize(val))
        if parts:
            pdf.set_font("Times", "", 9.5)
            pdf.cell(w, 4, " | ".join(parts), align="C",
                     new_x="LMARGIN", new_y="NEXT")

        # ── Horizontal rule below header ──
        pdf.ln(1.0)
        y = pdf.get_y()
        pdf.set_draw_color(0, 0, 0)
        pdf.set_line_width(0.4)
        pdf.line(m, y, m + w, y)
        pdf.ln(header_gap)

        def section_header(title: str):
            pdf.ln(section_gap)
            pdf.set_font("Times", "B", 11)
            pdf.cell(w, 4.5, title.upper(), new_x="LMARGIN", new_y="NEXT")
            ry = pdf.get_y()
            pdf.set_line_width(0.2)
            pdf.line(m, ry, m + w, ry)
            pdf.ln(post_header)

        def bullet(text: str, indent: float = 4):
            pdf.set_font("Times", "", body_size)
            pdf.set_x(pdf.l_margin + indent)
            pdf.cell(3, lh, "-")
            pdf.multi_cell(w - indent - 3, lh, sanitize(text),
                           new_x="LMARGIN", new_y="NEXT")

        # ── SUMMARY ──
        summary = tailored.get("summary", "")
        if summary:
            section_header("Summary")
            pdf.set_font("Times", "", body_size)
            pdf.multi_cell(w, lh, sanitize(summary),
                           new_x="LMARGIN", new_y="NEXT")

        # ── EXPERIENCE ──
        experience = tailored.get("experience", [])
        if experience:
            section_header("Experience")
            for i, exp in enumerate(experience):
                company = sanitize(exp.get("company", ""))
                location = sanitize(exp.get("location", ""))
                dates = sanitize(exp.get("dates", ""))
                company_line = f"{company}, {location}" if location else company

                pdf.set_font("Times", "B", 11)
                if dates:
                    pdf.set_font("Times", "", body_size)
                    date_w = pdf.get_string_width(dates) + 2
                    pdf.set_font("Times", "B", 11)
                    pdf.cell(w - date_w, 4.2, company_line)
                    pdf.set_font("Times", "", body_size)
                    pdf.cell(date_w, 4.2, dates, align="R",
                             new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.cell(w, 4.2, company_line,
                             new_x="LMARGIN", new_y="NEXT")

                title = sanitize(exp.get("title", ""))
                if title:
                    pdf.set_font("Times", "I", body_size)
                    pdf.cell(w, lh, title,
                             new_x="LMARGIN", new_y="NEXT")

                for b in exp.get("bullets", []):
                    bullet(b)
                if i < len(experience) - 1:
                    pdf.ln(entry_gap)

        # ── PROJECTS ──
        projects = tailored.get("projects", [])
        if projects:
            section_header("Projects")
            for i, proj in enumerate(projects):
                pdf.set_font("Times", "B", 11)
                pdf.cell(w, 4.5, sanitize(proj.get("name", "")),
                         new_x="LMARGIN", new_y="NEXT")

                for b in proj.get("bullets", []):
                    bullet(b)
                if i < len(projects) - 1:
                    pdf.ln(entry_gap)

        # ── TECHNICAL SKILLS ──
        skills = tailored.get("skills", {})
        if skills:
            section_header("Technical Skills")
            for category, skill_list in skills.items():
                if not skill_list:
                    continue
                pdf.set_font("Times", "B", body_size)
                cat_text = sanitize(f"{category}: ")
                cat_w = pdf.get_string_width(cat_text)
                pdf.cell(cat_w, lh_s, cat_text)
                pdf.set_font("Times", "", body_size)
                sk = ", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list)
                pdf.multi_cell(w - cat_w, lh_s, sanitize(sk),
                               new_x="LMARGIN", new_y="NEXT")
                pdf.ln(skill_gap)

        # ── EDUCATION ──
        education = tailored.get("education", [])
        if education:
            section_header("Education")
            for edu in education:
                inst = sanitize(edu.get("institution", ""))
                degree = sanitize(edu.get("degree", ""))
                dates = sanitize(edu.get("dates", ""))

                left_parts = [p for p in [inst, degree] if p]
                left_text = " | ".join(left_parts)

                pdf.set_font("Times", "B", body_size)
                if dates:
                    pdf.set_font("Times", "", body_size)
                    date_w = pdf.get_string_width(dates) + 2
                    pdf.set_font("Times", "B", body_size)
                    pdf.cell(w - date_w, lh_s, left_text)
                    pdf.set_font("Times", "", body_size)
                    pdf.cell(date_w, lh_s, dates, align="R",
                             new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.cell(w, lh_s, left_text,
                             new_x="LMARGIN", new_y="NEXT")

        content_h = pdf.get_y() - m
        if measure_only:
            return None, content_h
        return bytes(pdf.output()), content_h

    def generate_pdf(self, tailored: Dict[str, Any]) -> bytes:
        """Generate a single-page PDF that adaptively fills one A4 page.

        Two-pass approach:
          Pass 1: Render with minimum spacing to measure content height.
          Pass 2: Distribute remaining whitespace proportionally, then render final PDF.

        If content overflows even at minimum spacing, reduces font size and line heights.
        """
        avail = self._AVAIL_H
        slots = self._count_spacing_slots(tailored)

        # --- Pass 1: measure with minimum spacing ---
        _, min_height = self._render_pdf(
            tailored,
            section_gap=self._MIN_SECTION_GAP,
            entry_gap=self._MIN_ENTRY_GAP,
            post_header=self._MIN_POST_HEADER,
            header_gap=self._MIN_HEADER_GAP,
            skill_gap=self._MIN_SKILL_GAP,
            measure_only=True,
        )

        # --- Overflow protection: shrink if content exceeds one page at min spacing ---
        body_size = 10.0
        lh = 3.8
        lh_s = 3.6

        if min_height > avail:
            # Try smaller font first
            body_size = 9.5
            lh = 3.5
            lh_s = 3.3
            _, min_height = self._render_pdf(
                tailored,
                section_gap=self._MIN_SECTION_GAP,
                entry_gap=self._MIN_ENTRY_GAP,
                post_header=self._MIN_POST_HEADER,
                header_gap=self._MIN_HEADER_GAP,
                skill_gap=self._MIN_SKILL_GAP,
                body_size=body_size, lh=lh, lh_s=lh_s,
                measure_only=True,
            )

        if min_height > avail:
            # Even smaller
            body_size = 9.0
            lh = 3.2
            lh_s = 3.0
            _, min_height = self._render_pdf(
                tailored,
                section_gap=self._MIN_SECTION_GAP,
                entry_gap=self._MIN_ENTRY_GAP,
                post_header=self._MIN_POST_HEADER,
                header_gap=self._MIN_HEADER_GAP,
                skill_gap=self._MIN_SKILL_GAP,
                body_size=body_size, lh=lh, lh_s=lh_s,
                measure_only=True,
            )

        # --- Pass 2: distribute remaining whitespace ---
        remaining = max(0, avail - min_height)

        # Weight distribution: section gaps get 40%, entry gaps 30%,
        # post-header 15%, header gap 10%, skill gaps 5%
        n_sections = max(slots["sections"], 1)
        n_entry_gaps = max(slots["entry_gaps"], 1)
        n_skill_rows = max(slots["skill_rows"], 1)

        total_slots = (
            n_sections * 4      # section_gap weight
            + n_entry_gaps * 3  # entry_gap weight
            + n_sections * 1.5  # post_header weight
            + 1 * 1             # header_gap weight
            + n_skill_rows * 0.5  # skill_gap weight
        )

        if total_slots > 0 and remaining > 0:
            unit = remaining / total_slots
        else:
            unit = 0

        section_gap = min(self._MIN_SECTION_GAP + unit * 4, self._MAX_SECTION_GAP)
        entry_gap = min(self._MIN_ENTRY_GAP + unit * 3, self._MAX_ENTRY_GAP)
        post_header = min(self._MIN_POST_HEADER + unit * 1.5, self._MAX_POST_HEADER)
        header_gap = min(self._MIN_HEADER_GAP + unit * 1, self._MAX_HEADER_GAP)
        skill_gap = min(self._MIN_SKILL_GAP + unit * 0.5, self._MAX_SKILL_GAP)

        # --- Final render ---
        pdf_bytes, _ = self._render_pdf(
            tailored,
            section_gap=section_gap,
            entry_gap=entry_gap,
            post_header=post_header,
            header_gap=header_gap,
            skill_gap=skill_gap,
            body_size=body_size, lh=lh, lh_s=lh_s,
        )
        return pdf_bytes

    # ------------------------------------------------------------------
    # DOCX generation (python-docx)
    # ------------------------------------------------------------------

    def generate_docx(self, tailored: Dict[str, Any]) -> bytes:
        """Generate a single-page DOCX matching the LaTeX resume template."""
        from docx import Document
        from docx.shared import Pt, Inches, Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.oxml.ns import qn

        FONT = "Times New Roman"
        doc = Document()

        # A4 page, 0.4in margins
        for section in doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)

        # Set Normal style defaults
        style = doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = Pt(11)

        # Usable width for tab stops (A4 width - margins)
        page_w = Mm(210) - 2 * Inches(0.4)

        contact = tailored.get("contact", {})

        # ── Name (centered, bold, large) ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact.get("name", ""))
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = FONT
        p.paragraph_format.space_after = Pt(1)

        # ── Contact line (centered) ──
        parts = []
        for key in ("phone", "email", "linkedin", "github"):
            val = contact.get(key, "")
            if val:
                parts.append(val)
        if parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(parts))
            run.font.size = Pt(9)
            run.font.name = FONT
            p.paragraph_format.space_after = Pt(2)

        # ── HR line (bottom border on an empty paragraph) ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {qn("w:val"): "single", qn("w:sz"): "6",
             qn("w:space"): "1", qn("w:color"): "000000"},
        )
        pBdr.append(bottom)
        pPr.append(pBdr)

        def add_section_header(title: str):
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = FONT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {qn("w:val"): "single", qn("w:sz"): "4",
                 qn("w:space"): "1", qn("w:color"): "000000"},
            )
            pBdr.append(bottom)
            pPr.append(pBdr)

        def add_bullet(text: str):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = FONT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Pt(12)

        # ── SUMMARY ──
        summary = tailored.get("summary", "")
        if summary:
            add_section_header("Summary")
            p = doc.add_paragraph()
            run = p.add_run(summary)
            run.font.size = Pt(10)
            run.font.name = FONT
            p.paragraph_format.space_after = Pt(0)

        # ── EXPERIENCE ──
        experience = tailored.get("experience", [])
        if experience:
            add_section_header("Experience")
            for exp in experience:
                company = exp.get("company", "")
                location = exp.get("location", "")
                dates = exp.get("dates", "")
                company_line = f"{company}, {location}" if location else company

                # Company (bold) + Dates (right-aligned via tab stop)
                p = doc.add_paragraph()
                p.paragraph_format.tab_stops.add_tab_stop(
                    page_w, WD_TAB_ALIGNMENT.RIGHT)
                run = p.add_run(company_line)
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = FONT
                if dates:
                    run2 = p.add_run(f"\t{dates}")
                    run2.font.size = Pt(10)
                    run2.font.name = FONT
                p.paragraph_format.space_before = Pt(1.5)
                p.paragraph_format.space_after = Pt(0)

                # Job title (italic)
                title = exp.get("title", "")
                if title:
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run(title)
                    run2.italic = True
                    run2.font.size = Pt(10)
                    run2.font.name = FONT
                    p2.paragraph_format.space_after = Pt(0)

                for b in exp.get("bullets", []):
                    add_bullet(b)

        # ── PROJECTS ──
        projects = tailored.get("projects", [])
        if projects:
            add_section_header("Projects")
            for proj in projects:
                p = doc.add_paragraph()
                run = p.add_run(proj.get("name", ""))
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = FONT
                p.paragraph_format.space_before = Pt(1.5)
                p.paragraph_format.space_after = Pt(0)

                for b in proj.get("bullets", []):
                    add_bullet(b)

        # ── TECHNICAL SKILLS ──
        skills = tailored.get("skills", {})
        if skills:
            add_section_header("Technical Skills")
            for category, skill_list in skills.items():
                if not skill_list:
                    continue
                p = doc.add_paragraph()
                run_cat = p.add_run(f"{category}: ")
                run_cat.bold = True
                run_cat.font.size = Pt(10)
                run_cat.font.name = FONT
                sk = ", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list)
                run_sk = p.add_run(sk)
                run_sk.font.size = Pt(10)
                run_sk.font.name = FONT
                p.paragraph_format.space_after = Pt(0)

        # ── EDUCATION ──
        education = tailored.get("education", [])
        if education:
            add_section_header("Education")
            for edu in education:
                inst = edu.get("institution", "")
                degree = edu.get("degree", "")
                dates = edu.get("dates", "")

                left_parts = [x for x in [inst, degree] if x]
                left_text = " | ".join(left_parts)

                p = doc.add_paragraph()
                p.paragraph_format.tab_stops.add_tab_stop(
                    page_w, WD_TAB_ALIGNMENT.RIGHT)
                run = p.add_run(left_text)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT
                if dates:
                    run2 = p.add_run(f"\t{dates}")
                    run2.font.size = Pt(10)
                    run2.font.name = FONT
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(0)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Filename helper
    # ------------------------------------------------------------------

    @staticmethod
    def build_filename(tailored: Dict[str, Any], jd_analysis: Dict[str, Any], ext: str) -> str:
        """Build filename: Firstname_Lastname_JobTitle_YYYY-MM-DD.ext"""
        contact = tailored.get("contact", {})
        name = contact.get("name", "").strip()
        parts = name.split()
        first = parts[0] if parts else "Resume"
        last = parts[-1] if len(parts) > 1 else ""

        job_title = jd_analysis.get("job_title", "")

        def clean(s: str) -> str:
            s = re.sub(r"[^a-zA-Z0-9]+", "_", s).strip("_")
            return s[:30]

        date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        segments = [clean(first), clean(last), clean(job_title), date_str]
        segments = [s for s in segments if s]
        return "_".join(segments) + f".{ext}"
