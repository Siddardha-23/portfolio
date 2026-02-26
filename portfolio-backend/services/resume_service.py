"""
Resume Tailoring Service — JD extraction, resume tailoring, ATS scoring, document generation.

Each step is a standalone method called directly from the blueprint:
  1. Extract JD fields (Gemini)
  2. Tailor resume to JD (Gemini) — never fabricates, only rewords/reorders
  3. Compute ATS & AI screener scores (Gemini)
  4. Generate PDF (fpdf2) / DOCX (python-docx) matching the base resume format
"""
import io
import json
import logging
import re
from datetime import datetime, timezone
from typing import Any, Dict, Optional

from utils.db_connect import DBConnect

logger = logging.getLogger(__name__)

GEMINI_MODEL = "gemini-2.5-flash"


def _clean_json_response(text: str) -> str:
    """Strip markdown fences and trailing garbage from a Gemini response."""
    t = text.strip()
    # Remove ```json ... ``` wrappers
    if t.startswith("```"):
        t = re.sub(r"^```(?:json)?\s*\n?", "", t)
        t = re.sub(r"\n?```\s*$", "", t)
    # Trim anything after the last closing brace/bracket
    for end_char in ("}", "]"):
        idx = t.rfind(end_char)
        if idx != -1:
            t = t[: idx + 1]
            break
    return t.strip()


def _gemini_json(prompt: str, max_tokens: int = 8192, temperature: float = 0.3) -> dict:
    """Call Gemini and parse JSON response. Retries once with doubled tokens on truncation."""
    from services.chat_service import _get_client
    from google.genai import types

    client = _get_client()

    for attempt, tokens in enumerate([max_tokens, max_tokens * 2]):
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=temperature,
                max_output_tokens=tokens,
            ),
        )

        raw = (response.text or "").strip()
        if not raw:
            if attempt == 0:
                logger.warning("Gemini returned empty response, retrying with more tokens")
                continue
            raise ValueError("AI returned an empty response. Please try again.")

        # Try direct parse
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            pass

        # Clean and retry parse
        cleaned = _clean_json_response(raw)
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            # Likely truncated — retry with more tokens on first attempt
            if attempt == 0:
                logger.warning(f"JSON truncated (tokens={tokens}), retrying with {tokens * 2}")
                continue
            logger.error(f"Gemini invalid JSON after retry. Raw (first 500): {raw[:500]}")
            raise ValueError("AI returned an invalid response. Please try again.")


# ---------------------------------------------------------------------------
# Service class
# ---------------------------------------------------------------------------

class ResumeService:
    def __init__(self):
        db = DBConnect().get_db()
        self.user_resumes = db.user_resumes

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def get_base_resume(self) -> Optional[Dict[str, Any]]:
        resume = self.user_resumes.find_one({}, sort=[("parsed_at", -1)])
        if resume:
            resume.pop("_id", None)
        return resume

    # ------------------------------------------------------------------
    # Step 1 — Extract JD fields
    # ------------------------------------------------------------------

    def extract_jd(self, jd_text: str) -> Dict[str, Any]:
        prompt = (
            "You are an expert recruiter and ATS specialist.\n"
            "Extract structured information from the following job description.\n"
            "Return a JSON object with EXACTLY these keys:\n"
            '- "job_title": string (the specific role title)\n'
            '- "company": string (company name, or "Not specified")\n'
            '- "location": string (location or "Not specified")\n'
            '- "employment_type": string (Full-time/Part-time/Contract/Intern)\n'
            '- "required_skills": list of strings (must-have technical skills)\n'
            '- "preferred_skills": list of strings (nice-to-have skills)\n'
            '- "responsibilities": list of strings (key job duties, max 8)\n'
            '- "qualifications": list of strings (degree, years of experience, etc.)\n'
            '- "experience_years": string (e.g. "3-5 years" or "Entry level")\n'
            '- "industry": string (e.g. "Cloud Computing", "FinTech")\n'
            '- "keywords": list of strings (ATS-critical keywords from the JD)\n\n'
            f"=== JOB DESCRIPTION ===\n{jd_text[:8000]}"
        )
        return _gemini_json(prompt, max_tokens=4096)

    # ------------------------------------------------------------------
    # Step 2 — Tailor resume
    # ------------------------------------------------------------------

    def tailor_resume(self, raw_resume_text: str, jd_analysis: Dict[str, Any]) -> Dict[str, Any]:
        jd_json = json.dumps(jd_analysis, indent=2)
        # Build a keyword checklist string the AI must embed
        required = jd_analysis.get("required_skills", [])
        keywords = jd_analysis.get("keywords", [])
        keyword_list = ", ".join(dict.fromkeys(required + keywords))  # deduped, ordered

        prompt = (
            "You are a top-tier professional resume writer and ATS optimization expert.\n"
            "Given the candidate's ACTUAL resume text and a structured job description analysis,\n"
            "produce a COMPLETE tailored resume as a JSON object.\n\n"
            "YOUR #1 GOAL: Maximize ATS keyword match AND AI screener relevance score.\n"
            "The resume MUST contain these JD keywords/skills where the candidate has ANY related "
            f"experience: {keyword_list}\n\n"
            "ATS OPTIMIZATION RULES:\n"
            "1. Mirror the EXACT terminology from the JD (e.g. if JD says 'microservices', use 'microservices' not 'micro-services').\n"
            "2. Put the most JD-relevant skills FIRST in each category.\n"
            "3. Front-load each bullet with a strong action verb that matches the JD's language.\n"
            "4. Weave required keywords naturally into experience bullets and summary — don't just list them.\n"
            "5. The professional summary MUST mention the target role title and 3-4 top required skills.\n"
            "6. Reorder experience bullets so the most JD-relevant achievements appear first.\n"
            "7. For projects, emphasize aspects that directly relate to the JD requirements.\n\n"
            "INTEGRITY RULES:\n"
            "1. NEVER fabricate, invent, or add experience, projects, or job titles NOT in the resume.\n"
            "2. NEVER add major skills the candidate does not possess.\n"
            "3. You MAY add small, closely related skills (e.g., if they know Docker, you can add 'containerization').\n"
            "4. Keep ALL contact info and education dates EXACTLY as-is.\n"
            "5. Each experience bullet MUST include quantifiable impact where present in the original.\n"
            "6. Every field in the JSON MUST be a non-null string or array — never null.\n\n"
            "CONTENT & FORMAT RULES (the resume MUST fill ONE full page — A4, 0.4in margins, 10pt Times font):\n"
            "1. Summary: 2-3 sentences that fill 2-3 lines.\n"
            "2. Experience: Include ALL positions from the original resume. Each position gets 4-5 detailed "
            "bullets (each bullet should be 1-2 full lines of text, ~150-200 chars). Use strong action verbs "
            "and include metrics/numbers wherever possible.\n"
            "3. Projects: Include 1-3 projects with 3-4 detailed bullets each.\n"
            "4. Skills: 4-6 categories with comprehensive skill lists, most JD-relevant first.\n"
            "5. Education: Institution, degree, and dates.\n"
            "6. Do NOT include a 'certifications' section. Weave important certs into skills or summary.\n"
            "7. IMPORTANT: Generate ENOUGH content to fill a full page. Do NOT be overly brief. "
            "Each bullet should be substantive and detailed, not just a few words.\n\n"
            "Return a JSON object with EXACTLY this structure:\n"
            "{\n"
            '  "contact": {\n'
            '    "name": "Full Name",\n'
            '    "email": "email",\n'
            '    "phone": "phone",\n'
            '    "linkedin": "linkedin url or profile path",\n'
            '    "github": "github url or username"\n'
            "  },\n"
            '  "summary": "2-3 sentence professional summary tailored to the role",\n'
            '  "skills": {\n'
            '    "Category Name": ["Skill1", "Skill2", ...],\n'
            '    "Another Category": ["Skill3", "Skill4", ...]\n'
            "  },\n"
            '  "experience": [\n'
            "    {\n"
            '      "title": "Job Title",\n'
            '      "company": "Company Name",\n'
            '      "location": "City, State/Country",\n'
            '      "dates": "Start – End",\n'
            '      "type": "Full-time/Internship",\n'
            '      "bullets": ["Achievement 1...", "Achievement 2..."]\n'
            "    }\n"
            "  ],\n"
            '  "education": [\n'
            "    {\n"
            '      "degree": "Degree Name",\n'
            '      "institution": "University Name",\n'
            '      "location": "City, State",\n'
            '      "dates": "Start – End",\n'
            '      "gpa": "GPA if available",\n'
            '      "coursework": "Relevant coursework"\n'
            "    }\n"
            "  ],\n"
            '  "projects": [\n'
            "    {\n"
            '      "name": "Project Name",\n'
            '      "dates": "Date Range",\n'
            '      "bullets": ["Description 1...", "Description 2..."],\n'
            '      "tech": "Tech1, Tech2, Tech3"\n'
            "    }\n"
            "  ]\n"
            "}\n\n"
            f"=== ORIGINAL RESUME TEXT ===\n{raw_resume_text[:8000]}\n\n"
            f"=== JOB DESCRIPTION ANALYSIS ===\n{jd_json}"
        )
        return _gemini_json(prompt, max_tokens=8192, temperature=0.4)

    # ------------------------------------------------------------------
    # Step 3 — ATS & AI screener scoring
    # ------------------------------------------------------------------

    def compute_ats_scores(
        self, tailored: Dict[str, Any], jd_analysis: Dict[str, Any]
    ) -> Dict[str, Any]:
        tailored_json = json.dumps(tailored, indent=2)[:6000]
        jd_json = json.dumps(jd_analysis, indent=2)[:3000]

        # Build an explicit keyword checklist for the scorer
        required = jd_analysis.get("required_skills", [])
        keywords = jd_analysis.get("keywords", [])
        all_kw = list(dict.fromkeys(required + keywords))

        prompt = (
            "You are an ATS (Applicant Tracking System) and AI recruitment screener expert.\n"
            "Analyze the tailored resume against the job description.\n\n"
            "SCORING METHODOLOGY:\n"
            "1. keyword_match: For each keyword below, check if it appears VERBATIM in the resume. "
            f"Keywords to check: {', '.join(all_kw)}\n"
            "   Score = (keywords found / total keywords) * 100.\n"
            "2. skills_alignment: % of JD required_skills present in resume skills section.\n"
            "3. experience_relevance: How closely do job titles, industries, and responsibilities match.\n"
            "4. quantifiable_impact: Count bullets with specific numbers/metrics/percentages.\n"
            "5. format_score: ATS-friendly = no tables, no graphics, clear section headers, "
            "standard section names (Summary, Skills, Experience, Education).\n"
            "6. section_completeness: Has all of: Summary, Skills, Experience, Education, Certifications.\n"
            "7. Scanner scores: Each ATS parses differently — Workday favors exact keyword matches, "
            "Greenhouse weighs skills sections, Lever checks experience depth, "
            "iCIMS does strict keyword frequency, Taleo is format-sensitive, "
            "SmartRecruiters uses AI matching.\n"
            "8. AI screener: Simulates HireVue/Pymetrics-style screening for relevance, seniority fit, and tone.\n\n"
            "Be REALISTIC. If a required skill is missing, deduct heavily.\n"
            "All values MUST be integers 0-100. All arrays MUST be non-empty.\n\n"
            "Return a JSON object with EXACTLY this structure:\n"
            "{\n"
            '  "overall": integer,\n'
            '  "keyword_match": integer,\n'
            '  "skills_alignment": integer,\n'
            '  "experience_relevance": integer,\n'
            '  "quantifiable_impact": integer,\n'
            '  "format_score": integer,\n'
            '  "section_completeness": integer,\n'
            '  "scanners": { "workday": int, "greenhouse": int, "lever": int, '
            '"icims": int, "taleo": int, "smartrecruiters": int },\n'
            '  "ai_screener": { "overall": int, "relevance": int, '
            '"seniority_fit": int, "culture_fit": int },\n'
            '  "suggestions": ["actionable suggestion 1", "suggestion 2", "suggestion 3"],\n'
            '  "missing_keywords": ["keyword from JD not found in resume"],\n'
            '  "strengths": ["strong point 1", "strong point 2"]\n'
            "}\n\n"
            f"=== TAILORED RESUME ===\n{tailored_json}\n\n"
            f"=== JOB DESCRIPTION ANALYSIS ===\n{jd_json}"
        )
        return _gemini_json(prompt, max_tokens=4096, temperature=0.3)


    # ------------------------------------------------------------------
    # PDF generation (fpdf2)
    # ------------------------------------------------------------------

    @staticmethod
    def _sanitize_text(text) -> str:
        """Replace Unicode chars unsupported by built-in PDF fonts with latin-1 equivalents."""
        if not text:
            return ""
        text = str(text)
        replacements = {
            "\u2022": "-",   # bullet •
            "\u2013": "-",   # en-dash –
            "\u2014": "--",  # em-dash —
            "\u2018": "'",   # left single quote
            "\u2019": "'",   # right single quote / apostrophe
            "\u201c": '"',   # left double quote
            "\u201d": '"',   # right double quote
            "\u2026": "...", # ellipsis …
            "\u00a0": " ",   # non-breaking space
            "\u200b": "",    # zero-width space
        }
        for char, repl in replacements.items():
            text = text.replace(char, repl)
        # Fallback: strip any remaining non-latin-1 characters
        return text.encode("latin-1", errors="replace").decode("latin-1")

    def generate_pdf(self, tailored: Dict[str, Any]) -> bytes:
        """Generate a single-page PDF matching the LaTeX resume template.

        Format: A4, 0.4in margins, Times font, tight spacing.
        Section order: Summary → Experience → Projects → Technical Skills → Education.
        """
        from fpdf import FPDF
        sanitize = self._sanitize_text

        contact = tailored.get("contact", {})
        name = sanitize(contact.get("name", "Resume"))

        pdf = FPDF(format="A4")
        pdf.set_auto_page_break(auto=True, margin=10)
        pdf.add_page()

        m = 10.16  # 0.4 inch margins in mm
        pdf.set_margins(m, m, m)
        pdf.set_y(m)
        w = pdf.w - 2 * m  # usable width ~189.7mm

        LH = 4.0   # line height for 10pt body text
        LH_S = 3.6  # tighter line height for skills/education

        # ── Name (centered, bold, large) ──
        pdf.set_font("Times", "B", 22)
        pdf.cell(w, 8, name, align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(0.5)

        # ── Contact line ──
        parts = []
        for key in ("phone", "email", "linkedin", "github"):
            val = contact.get(key, "")
            if val:
                parts.append(sanitize(val))
        if parts:
            pdf.set_font("Times", "", 9.5)
            pdf.cell(w, 4, " | ".join(parts), align="C",
                     new_x="LMARGIN", new_y="NEXT")

        # ── Horizontal rule below header ──
        pdf.ln(1.5)
        y = pdf.get_y()
        pdf.set_draw_color(0, 0, 0)
        pdf.set_line_width(0.4)
        pdf.line(m, y, m + w, y)
        pdf.ln(2)

        def section_header(title: str):
            pdf.ln(1.5)
            pdf.set_font("Times", "B", 11)
            pdf.cell(w, 5, title.upper(), new_x="LMARGIN", new_y="NEXT")
            y = pdf.get_y()
            pdf.set_line_width(0.2)
            pdf.line(m, y, m + w, y)
            pdf.ln(1.5)

        def bullet(text: str, indent: float = 4):
            pdf.set_font("Times", "", 10)
            pdf.set_x(pdf.l_margin + indent)
            pdf.cell(3, LH, "-")
            pdf.multi_cell(w - indent - 3, LH, sanitize(text),
                           new_x="LMARGIN", new_y="NEXT")

        # ── SUMMARY ──
        summary = tailored.get("summary", "")
        if summary:
            section_header("Summary")
            pdf.set_font("Times", "", 10)
            pdf.multi_cell(w, LH, sanitize(summary),
                           new_x="LMARGIN", new_y="NEXT")

        # ── EXPERIENCE ──
        experience = tailored.get("experience", [])
        if experience:
            section_header("Experience")
            for i, exp in enumerate(experience):
                company = sanitize(exp.get("company", ""))
                location = sanitize(exp.get("location", ""))
                dates = sanitize(exp.get("dates", ""))
                company_line = f"{company}, {location}" if location else company

                # Company (bold left) + Dates (right)
                pdf.set_font("Times", "B", 11)
                if dates:
                    pdf.set_font("Times", "", 10)
                    date_w = pdf.get_string_width(dates) + 2
                    pdf.set_font("Times", "B", 11)
                    pdf.cell(w - date_w, 4.5, company_line)
                    pdf.set_font("Times", "", 10)
                    pdf.cell(date_w, 4.5, dates, align="R",
                             new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.cell(w, 4.5, company_line,
                             new_x="LMARGIN", new_y="NEXT")

                # Job title (italic)
                title = sanitize(exp.get("title", ""))
                if title:
                    pdf.set_font("Times", "I", 10)
                    pdf.cell(w, LH, title,
                             new_x="LMARGIN", new_y="NEXT")

                for b in exp.get("bullets", []):
                    bullet(b)
                if i < len(experience) - 1:
                    pdf.ln(1.5)

        # ── PROJECTS ──
        projects = tailored.get("projects", [])
        if projects:
            section_header("Projects")
            for i, proj in enumerate(projects):
                pdf.set_font("Times", "B", 11)
                pdf.cell(w, 4.5, sanitize(proj.get("name", "")),
                         new_x="LMARGIN", new_y="NEXT")

                for b in proj.get("bullets", []):
                    bullet(b)
                if i < len(projects) - 1:
                    pdf.ln(1.5)

        # ── TECHNICAL SKILLS ──
        skills = tailored.get("skills", {})
        if skills:
            section_header("Technical Skills")
            for category, skill_list in skills.items():
                if not skill_list:
                    continue
                pdf.set_font("Times", "B", 10)
                cat_text = sanitize(f"{category}: ")
                cat_w = pdf.get_string_width(cat_text)
                pdf.cell(cat_w, LH_S, cat_text)
                pdf.set_font("Times", "", 10)
                sk = ", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list)
                pdf.multi_cell(w - cat_w, LH_S, sanitize(sk),
                               new_x="LMARGIN", new_y="NEXT")
                pdf.ln(0.3)

        # ── EDUCATION ──
        education = tailored.get("education", [])
        if education:
            section_header("Education")
            for edu in education:
                inst = sanitize(edu.get("institution", ""))
                degree = sanitize(edu.get("degree", ""))
                dates = sanitize(edu.get("dates", ""))

                # Bold "Institution | Degree" left + dates right
                left_parts = [p for p in [inst, degree] if p]
                left_text = " | ".join(left_parts)

                pdf.set_font("Times", "B", 10)
                if dates:
                    pdf.set_font("Times", "", 10)
                    date_w = pdf.get_string_width(dates) + 2
                    pdf.set_font("Times", "B", 10)
                    pdf.cell(w - date_w, LH_S, left_text)
                    pdf.set_font("Times", "", 10)
                    pdf.cell(date_w, LH_S, dates, align="R",
                             new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.cell(w, LH_S, left_text,
                             new_x="LMARGIN", new_y="NEXT")

        return pdf.output()

    # ------------------------------------------------------------------
    # DOCX generation (python-docx)
    # ------------------------------------------------------------------

    def generate_docx(self, tailored: Dict[str, Any]) -> bytes:
        """Generate a single-page DOCX matching the LaTeX resume template.

        Format: A4, 0.4in margins, Times New Roman, tight spacing.
        Section order: Summary → Experience → Projects → Technical Skills → Education.
        """
        from docx import Document
        from docx.shared import Pt, Inches, Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.oxml.ns import qn

        FONT = "Times New Roman"
        doc = Document()

        # A4 page, 0.4in margins
        for section in doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)

        # Set Normal style defaults
        style = doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = Pt(12)

        # Usable width for tab stops (A4 width - margins)
        page_w = Mm(210) - 2 * Inches(0.4)

        contact = tailored.get("contact", {})

        # ── Name (centered, bold, large) ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact.get("name", ""))
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = FONT
        p.paragraph_format.space_after = Pt(1)

        # ── Contact line (centered) ──
        parts = []
        for key in ("phone", "email", "linkedin", "github"):
            val = contact.get(key, "")
            if val:
                parts.append(val)
        if parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(parts))
            run.font.size = Pt(9)
            run.font.name = FONT
            p.paragraph_format.space_after = Pt(2)

        # ── HR line (bottom border on an empty paragraph) ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {qn("w:val"): "single", qn("w:sz"): "6",
             qn("w:space"): "1", qn("w:color"): "000000"},
        )
        pBdr.append(bottom)
        pPr.append(pBdr)

        def add_section_header(title: str):
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = FONT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {qn("w:val"): "single", qn("w:sz"): "4",
                 qn("w:space"): "1", qn("w:color"): "000000"},
            )
            pBdr.append(bottom)
            pPr.append(pBdr)

        def add_bullet(text: str):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = FONT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Pt(14)

        # ── SUMMARY ──
        summary = tailored.get("summary", "")
        if summary:
            add_section_header("Summary")
            p = doc.add_paragraph()
            run = p.add_run(summary)
            run.font.size = Pt(10)
            run.font.name = FONT
            p.paragraph_format.space_after = Pt(1)

        # ── EXPERIENCE ──
        experience = tailored.get("experience", [])
        if experience:
            add_section_header("Experience")
            for exp in experience:
                company = exp.get("company", "")
                location = exp.get("location", "")
                dates = exp.get("dates", "")
                company_line = f"{company}, {location}" if location else company

                # Company (bold) + Dates (right-aligned via tab stop)
                p = doc.add_paragraph()
                p.paragraph_format.tab_stops.add_tab_stop(
                    page_w, WD_TAB_ALIGNMENT.RIGHT)
                run = p.add_run(company_line)
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = FONT
                if dates:
                    run2 = p.add_run(f"\t{dates}")
                    run2.font.size = Pt(10)
                    run2.font.name = FONT
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(0)

                # Job title (italic)
                title = exp.get("title", "")
                if title:
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run(title)
                    run2.italic = True
                    run2.font.size = Pt(10)
                    run2.font.name = FONT
                    p2.paragraph_format.space_after = Pt(0)

                for b in exp.get("bullets", []):
                    add_bullet(b)

        # ── PROJECTS ──
        projects = tailored.get("projects", [])
        if projects:
            add_section_header("Projects")
            for proj in projects:
                p = doc.add_paragraph()
                run = p.add_run(proj.get("name", ""))
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = FONT
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(0)

                for b in proj.get("bullets", []):
                    add_bullet(b)

        # ── TECHNICAL SKILLS ──
        skills = tailored.get("skills", {})
        if skills:
            add_section_header("Technical Skills")
            for category, skill_list in skills.items():
                if not skill_list:
                    continue
                p = doc.add_paragraph()
                run_cat = p.add_run(f"{category}: ")
                run_cat.bold = True
                run_cat.font.size = Pt(10)
                run_cat.font.name = FONT
                sk = ", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list)
                run_sk = p.add_run(sk)
                run_sk.font.size = Pt(10)
                run_sk.font.name = FONT
                p.paragraph_format.space_after = Pt(0)

        # ── EDUCATION ──
        education = tailored.get("education", [])
        if education:
            add_section_header("Education")
            for edu in education:
                inst = edu.get("institution", "")
                degree = edu.get("degree", "")
                dates = edu.get("dates", "")

                left_parts = [x for x in [inst, degree] if x]
                left_text = " | ".join(left_parts)

                p = doc.add_paragraph()
                p.paragraph_format.tab_stops.add_tab_stop(
                    page_w, WD_TAB_ALIGNMENT.RIGHT)
                run = p.add_run(left_text)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT
                if dates:
                    run2 = p.add_run(f"\t{dates}")
                    run2.font.size = Pt(10)
                    run2.font.name = FONT
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(0)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Filename helper
    # ------------------------------------------------------------------

    @staticmethod
    def build_filename(tailored: Dict[str, Any], jd_analysis: Dict[str, Any], ext: str) -> str:
        """Build filename: Firstname_Lastname_JobTitle_YYYY-MM-DD.ext"""
        contact = tailored.get("contact", {})
        name = contact.get("name", "").strip()
        parts = name.split()
        first = parts[0] if parts else "Resume"
        last = parts[-1] if len(parts) > 1 else ""

        job_title = jd_analysis.get("job_title", "")

        def clean(s: str) -> str:
            # Replace spaces/special chars with underscores, collapse multiples
            s = re.sub(r"[^a-zA-Z0-9]+", "_", s).strip("_")
            return s[:30]

        date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        segments = [clean(first), clean(last), clean(job_title), date_str]
        segments = [s for s in segments if s]
        return "_".join(segments) + f".{ext}"


# Singleton
_resume_service: Optional[ResumeService] = None


def get_resume_service() -> ResumeService:
    global _resume_service
    if _resume_service is None:
        _resume_service = ResumeService()
    return _resume_service
